# -*- coding: utf-8 -*-
"""Gera o Kit de Prospecção de Marcas do Augusto Felipe (.docx)."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---- paleta ----
DARK   = RGBColor(0x1A, 0x1A, 0x1A)
ACCENT = RGBColor(0x2E, 0x9B, 0xFF)   # azul
GREEN  = RGBColor(0x1E, 0xA0, 0x5A)
GRAY   = RGBColor(0x66, 0x66, 0x66)
LIGHT  = RGBColor(0x8A, 0x8A, 0x8A)

doc = Document()

# fonte padrão
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10.5)
style.font.color.rgb = DARK
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.12

def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)

def set_cell_font(cell, size=9.5, bold=False, color=DARK, white=False):
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.size = Pt(size)
            r.font.bold = bold
            r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF) if white else color

def h1(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(20); r.font.bold = True; r.font.color.rgb = DARK
    p.paragraph_format.space_after = Pt(2)
    return p

def kicker(text):
    p = doc.add_paragraph()
    r = p.add_run(text.upper())
    r.font.size = Pt(9); r.font.bold = True; r.font.color.rgb = ACCENT
    r.font.name = "Calibri"
    # letter spacing
    p.paragraph_format.space_after = Pt(10)
    return p

def h2(text, num=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    if num:
        r0 = p.add_run(f"{num}  ")
        r0.font.size = Pt(15); r0.font.bold = True; r0.font.color.rgb = ACCENT
    r = p.add_run(text)
    r.font.size = Pt(15); r.font.bold = True; r.font.color.rgb = DARK
    return p

def h3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.font.size = Pt(12); r.font.bold = True; r.font.color.rgb = DARK
    return p

def body(text, color=DARK):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.color.rgb = color
    return p

def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.font.bold = True
    p.add_run(text)
    return p

def small(text, color=GRAY, italic=True):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(9); r.font.italic = italic; r.font.color.rgb = color
    return p

def hr():
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1"); bottom.set(qn("w:color"), "DDDDDD")
    pbdr.append(bottom); pPr.append(pbdr)
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(6)

def callout(title, text, fill="EAF4FF", accent=ACCENT):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.autofit = True
    cell = tbl.rows[0].cells[0]
    shade(cell, fill)
    p = cell.paragraphs[0]
    rt = p.add_run(title)
    rt.font.bold = True; rt.font.size = Pt(10.5); rt.font.color.rgb = accent
    p2 = cell.add_paragraph()
    r2 = p2.add_run(text)
    r2.font.size = Pt(10)
    cell.paragraphs[0].paragraph_format.space_after = Pt(3)
    # margins
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for tag in ("top","start","bottom","end"):
        m = OxmlElement(f"w:{tag}")
        m.set(qn("w:w"), "120"); m.set(qn("w:type"), "dxa")
        tcMar.append(m)
    tcPr.append(tcMar)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return tbl

def template_box(text):
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.rows[0].cells[0]
    shade(cell, "F5F5F5")
    # left border accent
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    left = OxmlElement("w:left")
    left.set(qn("w:val"),"single"); left.set(qn("w:sz"),"18")
    left.set(qn("w:space"),"0"); left.set(qn("w:color"),"2E9BFF")
    borders.append(left); tcPr.append(borders)
    tcMar = OxmlElement("w:tcMar")
    for tag in ("top","start","bottom","end"):
        m = OxmlElement(f"w:{tag}")
        m.set(qn("w:w"), "160" if tag in ("start","end") else "120"); m.set(qn("w:type"), "dxa")
        tcMar.append(m)
    tcPr.append(tcMar)
    first = True
    for line in text.split("\n"):
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        r = p.add_run(line)
        r.font.size = Pt(10); r.font.name = "Calibri"
        p.paragraph_format.space_after = Pt(2)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return tbl

def ref_table(hl, hr, pairs, left_w=2.1):
    """Tabela de 2 colunas com cabeçalho escuro. left_w em polegadas."""
    t = doc.add_table(rows=1, cols=2)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.autofit = False
    hcells = t.rows[0].cells
    hcells[0].text = hl; hcells[1].text = hr
    for c in hcells:
        shade(c, "1A1A1A"); set_cell_font(c, 9, True, white=True)
    for i,(a,b) in enumerate(pairs):
        cells = t.add_row().cells
        cells[0].text = a; cells[1].text = b
        shade(cells[0], "F0F4F8"); shade(cells[1], "FFFFFF" if i%2 else "FAFAFA")
        set_cell_font(cells[0], 9, bold=True)
        set_cell_font(cells[1], 9)
    right_w = max(2.2, 6.3 - left_w)
    for row in t.rows:
        row.cells[0].width = Inches(left_w)
        row.cells[1].width = Inches(right_w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t

# =====================================================================
# CAPA
# =====================================================================
kicker("Kit de Prospecção de Marcas")
h1("Augusto Felipe · Fi de Vidraceiro")
p = doc.add_paragraph()
r = p.add_run("Guia estratégico + roteiro + rate card + diretório de contatos + DMs prontas")
r.font.size = Pt(12); r.font.color.rgb = GRAY
p = doc.add_paragraph()
r = p.add_run("@fidevidraceiro  ·  1.2M seguidores  ·  59M de alcance/mês")
r.font.size = Pt(10); r.font.bold = True; r.font.color.rgb = ACCENT
small("Baseado em dados e estudos de caso de prospecção creator→marca de 2026.  ·  Uso interno.", GRAY)
hr()

# =====================================================================
# 1. O DIAGNÓSTICO
# =====================================================================
h2("O diagnóstico: por que e-mail frio não responde", "1")
body("O e-mail frio como PRIMEIRO contato morreu. Os dados de 2026 são claros — e confirmam o que você já sentiu na prática:")

# tabela de taxas
tbl = doc.add_table(rows=6, cols=2)
tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
tbl.style = "Table Grid"
rows = [
    ("Origem do contato", "Taxa de resposta 2026"),
    ("E-mail frio (B2B)", "1–5%  (era ~7% há 2 anos)"),
    ("DM Instagram direcionado", "24–32%  (frio 10–25% / morno 30–50%)"),
    ("LinkedIn InMail", "18–25%"),
    ("Inbound (a marca vem até você)", "60–80%"),
    ("", ""),
]
for i,(a,b) in enumerate(rows[:5]):
    tbl.rows[i].cells[0].text = a
    tbl.rows[i].cells[1].text = b
    header = (i==0)
    for j,cell in enumerate(tbl.rows[i].cells):
        if header:
            shade(cell, "1A1A1A"); set_cell_font(cell, 9.5, True, white=True)
        else:
            shade(cell, "FFFFFF" if i%2 else "F5F9FF")
            set_cell_font(cell, 9.5, bold=(j==1))
# remove extra row
tbl._tbl.remove(tbl.rows[5]._tr)
doc.add_paragraph().paragraph_format.space_after = Pt(2)

callout("O DM bate o e-mail frio em 5–6x.",
        "A conversa antiga de 'achar o e-mail e disparar proposta' é a de menor conversão que existe hoje. "
        "O perfil do Augusto é a arma: quem recebe o DM já vê 1.2M de seguidores e a qualidade do conteúdo no mesmo clique.")

h3("A nuance que resolve a estratégia (o 'conflito' dos dados)")
body("Você vai ver fontes dizendo 'e-mail ganha' e outras 'DM ganha'. As duas estão certas — sobre momentos diferentes:")
bullet("— o DM ganha, disparado. E-mail pra quem nunca ouviu falar de você é ignorado.", "Para ABRIR (frio): ")
bullet("— aí o e-mail ganha, porque dá espaço pra avaliar. Mas só DEPOIS que a relação esquentou. Mandar proposta de 5 dígitos num DM frio sinaliza inexperiência.", "Para FECHAR (proposta cheia): ")
body("Conclusão: DM abre, e-mail formaliza. A ordem é o que importa.", GRAY)

# =====================================================================
# 2. O QUE FUNCIONA
# =====================================================================
h2("O que os estudos de caso de 2026 mostram", "2")

h3("a) Aquecer 1–2 semanas ANTES de pitchar")
body("Antes de qualquer pedido: responder Stories da marca (de verdade, não só curtir), comentar posts com algo útil, e marcar a marca em conteúdo orgânico usando o produto. Quando o DM chega, o Augusto não é um estranho.")

h3("b) Personalização é o divisor de águas")
bullet("Mencionar um post das ÚLTIMAS 2 SEMANAS = +8–10% de resposta.")
bullet("Proposta customizada = 3x mais aceitação vs. genérica.")
bullet("76% das marcas dizem que o media kit influencia direto a decisão.")
small("Os percentuais acima são de estudos de mercado 2026 — servem de norte interno. Não cite esses números numa conversa com a marca; use os SEUS números (1.2M, 59M, views dos Reels).")

h3("c) Engajamento > seguidores")
body("As marcas em 2026 priorizam taxa de engajamento. 30K engajados batem 300K passivos. Criadores de 10K–100K estão fechando deals de 5 dígitos — o Augusto, com 1.2M e engajamento real, joga num patamar acima. E tem um diferencial que quase nenhum criador tem: transformar produto em arte de vidro. O argumento é o ENGAJAMENTO e esse formato único, não só o número.")

h3("d) Estrutura do DM — modelo INTEREST (máx. 150 palavras)")
for a,b in [
    ("Identificação", "mostra que sabe quem eles são"),
    ("Notice", "cita algo recente e específico (post das últimas 2 semanas)"),
    ("Tie-in", "por que faz sentido pra eles"),
    ("Explain", "a proposta, de forma clara"),
    ("Result", "qual resultado a marca ganha"),
    ("Ease", "UM próximo passo simples (1 link só)"),
]:
    p = doc.add_paragraph(style="List Bullet")
    rr = p.add_run(f"{a} — "); rr.font.bold = True
    p.add_run(b)
small("Acima de 150 palavras = -30% de resposta. Máx. 1–2 links, sem CAIXA ALTA (cai em spam).")

h3("e) Timing e cadência (comprovado)")
bullet("Enviar terça a quinta, 10h–14h. Evitar segunda e sexta.", "Melhor janela: ")
bullet("espera 3–5 dias pro 1º, depois 7–10 pro 2º. PARA no 3º toque do mesmo canal — o 4º soa desesperado.", "Follow-up no DM: ")
bullet("a maioria dos fechamentos leva 5–7 pontos de contato somando canais. Persistência é multicanal, não repetição.", "No total: ")

# =====================================================================
# 3. O ROTEIRO
# =====================================================================
h2("O roteiro do Augusto (passo a passo)", "3")
callout("A regra de ouro", "Cada canal tem UM trabalho: o Instagram ABRE a porta, o LinkedIn IDENTIFICA o decisor, o e-mail FORMALIZA com o media kit.", "EAF4FF")

steps = [
    ("Semana 0–1 · Aquecer", "Escolher 10–15 marcas-alvo e aquecer cada uma: responder Stories, comentar, marcar em conteúdo usando o produto. Sem pedir nada ainda."),
    ("Semana 1 · DM (pelo @fidevidraceiro)", "Mensagem no modelo INTEREST. O objetivo NÃO é mandar proposta — é descobrir o decisor e gerar interesse. Pergunta quem cuida das parcerias."),
    ("LinkedIn (em paralelo, marcas grandes)", "Buscar o decisor pelo cargo (Marketing / Mídia / Parcerias / Social / Branded Content). Conectar com nota curta. Em empresa média/regional, mirar o dono ou diretor de marketing."),
    ("Depois do 'pode mandar' · E-mail", "Aí sim: e-mail com o media kit + UMA ideia concreta pra aquela marca (não genérico). A proposta BYD é a base — troca o contexto."),
    ("Follow-up", "3–5 dias depois, um lembrete educado. Somando canais, insiste até ~5–7 toques. Sem responder após isso, deixa esfriar e retoma em 30–60 dias."),
]
for i,(t,d) in enumerate(steps, 1):
    h3(f"Passo {i} — {t}")
    body(d)

callout("A jogada de longo prazo: virar INBOUND (60–80%)",
        "Conteúdo que posiciona o Augusto como 'o criador que a marca X deveria contratar' faz a marca vir até ele. "
        "Ex.: vídeos usando produtos de marcas-sonho de forma criativa, marcando elas. É o canal de maior conversão que existe.",
        "EFFaF0", GREEN)

# =====================================================================
# 4. QUEM MIRAR
# =====================================================================
h2("Quem mirar dentro da empresa & quais marcas", "4")
h3("O decisor certo, por porte")
bullet("Gerente/Coord. de Marketing, Social Media, Branded Content ou Influencer Marketing (às vezes via agência).", "Grande (tipo BYD): ")
bullet("Diretor de Marketing ou o próprio dono/sócio.", "Média/regional: ")
bullet("Dono direto.", "Local/pequena: ")

callout("A lente é a ARTE — não o material que ele usa",
        "O ativo do Augusto não é 'vidro pra marca de construção'. É transformar QUALQUER produto em obra. A BYD prova: é montadora, não construção. "
        "A pergunta que qualifica uma marca é uma só: \"o que dá pra transformar em obra pra ELES?\" Se a resposta é forte, é alvo — do perfume ao energético.",
        "EAF4FF", ACCENT)

h3("Como priorizar as marcas")
bullet("A marca já rasga verba com influenciador? Se patrocina criador toda semana, o orçamento existe e a decisão é rápida. Este é o filtro nº 1.", "Investe pesado em criador: ")
bullet("Marca em alta / momento de mídia (lançamento, campanha viral, aniversário). Ex.: Cimed, Stanley, Nubank. A arte vira o 'momento' da campanha.", "Tem hype agora: ")
bullet("Produtos com forma icônica (frasco, garrafa, lata, tênis, carro) rendem uma peça de vidro óbvia. Quanto mais reconhecível o objeto, mais viral a obra.", "Rende arte forte: ")
bullet("1–2 marcas-sonho (BYD, Coca) + várias que já fazem ação com criador e fecham rápido pra virar case.", "Mix: ")
callout("Prova social que o Augusto já tem",
        "Marcas que já confiam no trabalho — use isso em TODA abordagem: TekBond, Cebrace, Bombril, Temu, Vonixx, Spotify e Call of Duty. "
        "Reels de alcance comprovado: 8M · 6,3M · 5,5M views.", "F5F5F5", DARK)

h3("12 marcas-alvo sugeridas — de qualquer nicho, com a ideia de arte")
small("Critério: hype + verba de influenciador comprovada. O que muda por marca não é o nicho — é o que o Augusto transforma em obra.")
ref_table("Marca (setor)", "A ideia de arte + por que agora", [
    ("Cimed (farma)", "Réplica gigante em vidro do Carmed / do produto do momento. Marca-símbolo de marketing viral no Brasil — arte é o próximo nível."),
    ("Stanley (lifestyle)", "O copo Stanley recriado como obra de vidro — objeto de desejo do momento, hype absurdo, público jovem."),
    ("Natura / O Boticário", "Frasco de perfume gigante esculpido em vidro. Perfume É vidro: encaixe perfeito e escada natural até o Full Art Experience."),
    ("Coca-Cola (bebida)", "A garrafa contour em vidro artístico — o objeto mais icônico do mundo, feito à mão. Verba de branded content enorme."),
    ("Red Bull (energético)", "Arte em vidro com as asas/lata. Red Bull vive de patrocinar o inusitado — é o parceiro ideal pra uma obra viral."),
    ("Heineken / Ambev (cerveja)", "Garrafa/estrela em vidro verde. Investem pesado em criador, evento e ativação de balada."),
    ("Nubank (fintech)", "Peça de vidro roxa — o roxo é a marca. Cartão gigante ou o 'Nu' esculpido. Fintech que adora campanha diferente e tem verba."),
    ("iFood (delivery/tech)", "Ativação divertida em vidro com a identidade vermelha. Marca jovem, meme-friendly, fecha muito com criador."),
    ("Havaianas (moda BR)", "Chinelo/verão traduzido em arte de vidro. Ícone brasileiro que ama collab de criador nacional."),
    ("Samsung / Motorola (tech)", "Tela/celular esculpido em vidro — o gancho é literal (Gorilla Glass). Tech investe alto em criador."),
    ("McDonald's / BK (fast food)", "O ícone (batata, casquinha) virando obra de vidro. QSR investe muito em conteúdo de criador e OOH."),
    ("GWM / Caoa Chery (automotivo)", "Reaproveita 100% o case BYD: o carro virando obra. Setor em plena expansão no Brasil, brigando por atenção."),
], left_w=2.3)
small("Isto é ponto de partida. Regra pra achar outras: marca com hype + que já posta com influenciador toda semana = alvo, seja qual for o nicho.")

h3("Melhor porta de entrada por marca (resumo)")
body("O contato mais aderente pra abrir cada marca. A lista COMPLETA — 3 a 5 opções por marca, com LinkedIn/Instagram e as agências de influência — está no §11 (Diretório de contatos por marca). Cargos giram rápido: confirme no LinkedIn antes de abordar.")
ref_table("Marca", "Por onde começar (top pick)", [
    ("Cimed", "Toguro — Head de Comunicação, é criador (IG @toguro) · Camille Lau — CMO. Marca ama meme/collab."),
    ("Stanley", "Patrícia Esteves — Dir. Executiva de Mkt LatAm (LinkedIn /paesteves) · Babi Coelho — Ger. Mkt BR."),
    ("Natura / Boticário", "Boticário: Carolina Carrasco — Dir. Branding · Renata Gomide (IG @reagomide) · ag. CoCreators. Natura: Tatiana Ponce — CMO · prog. Criadores da Beleza (YouPix)."),
    ("Coca-Cola", "Raquel Ribeiro — Head de Marketing BR. Influência via agência BR Media Group. Raphael Taborda — social operacional."),
    ("Red Bull", "Buscar \"Culture Marketing Manager Red Bull Brasil\" (quem lida com criadores) · via agência Publicis."),
    ("Heineken / Ambev", "Heineken: Fernanda Saboya — Consumer Connections/influência (LinkedIn /fsaboya). Ambev: Mariana Dedivitis (Corona) · agência Draftline/Soko · InPress (PR)."),
    ("Nubank", "Raquel Forastieri — Brand Mkt Manager (cuida de influência, LinkedIn /rquelforastieri) · Juliana Roschel — CMO. Influência é in-house."),
    ("iFood", "Marina Moreira — Head de Social & Influência (melhor porta, LinkedIn /marina-moreira-97520696) · Renata Lamarco — Diretora de Mkt."),
    ("Havaianas", "Karen Bartels — a.gente, agência de influência da marca (via GALERIA.ag) · Juliana Soncini — Dir. Mkt LATAM."),
    ("Samsung / Motorola", "Samsung: agência gen_c cuida da influência Mobile (Cristiane Vieira) · Mario Sousa — Dir. Sênior Mobile. Motorola: Andrea Brandi — Sr. Brand Mkt."),
    ("McDonald's / BK", "McDonald's: Guilherme Coe — Ger. Mkt · Ilca Sierra — CMO. BK: Aymê Alves — redes sociais (Zamp) · Pedro Barbosa — CMO."),
    ("GWM / Caoa Chery", "GWM: Lucas Coelho — CMO (LinkedIn /lucas-a-coelho) · Magda Lima — Ger. Mkt. Caoa Chery: Jan Telecki — Dir. Mkt · agência Baila."),
], left_w=1.7)

doc.add_page_break()

# =====================================================================
# 5. TEMPLATES
# =====================================================================
kicker("Parte prática")
h1("Templates prontos")
body("Preenchidos com os números reais do media kit. Personalize o [colchete] em cada envio — genérico não converte.", GRAY)

h2("Template A — DM de abertura (Instagram)", "A")
small("Enviado pelo @fidevidraceiro. Modelo INTEREST, < 150 palavras. NÃO manda proposta aqui.")
template_box(
"Fala, [nome/marca]! Sou o Augusto, o Fi de Vidraceiro (@fidevidraceiro).\n"
"\n"
"Vi o post de vocês sobre [algo específico das últimas 2 semanas] e curti demais.\n"
"\n"
"Trabalho transformando vidro em arte e mostrando os bastidores da construção civil pra uma comunidade de 1.2M de seguidores (59M de alcance/mês). Tenho uma ideia de conteúdo com [produto de vocês] que faz muito sentido pro público de vocês.\n"
"\n"
"Quem cuida das parcerias/marketing aí? Posso mandar meu media kit direto pra pessoa certa. 🙏")

h2("Template B — Convite no LinkedIn", "B")
small("Para chegar no decisor de marcas grandes. Nota curta junto do convite.")
template_box(
"Oi [nome], tudo bem? Sou o Augusto Felipe, criador do 'Fi de Vidraceiro' (@fidevidraceiro, 1.2M seguidores).\n"
"\n"
"Tenho uma ideia de collab que combina muito com a [marca] e queria te apresentar. Posso te mandar meu media kit?")

h2("Template C — E-mail de proposta (após o 'pode mandar')", "C")
small("Só depois que a relação esquentou. Assunto curto + ideia concreta + CTA único.")
template_box(
"Assunto: Ideia de collab: Fi de Vidraceiro × [Marca]\n"
"\n"
"Oi [nome], como combinei no [Instagram/LinkedIn], segue minha ideia.\n"
"\n"
"Sou o Augusto, o Fi de Vidraceiro. Transformo vidro em arte e mostro os bastidores da construção civil — um universo que impacta 1.2M de seguidores e 59M de pessoas por mês (Instagram 730k · TikTok 430k · Kwai 110k). Marcas como TekBond, Cebrace, Bombril, Vonixx e Spotify já confiaram no trabalho.\n"
"\n"
"A ideia pra [marca]: [1 frase concreta — ex.: um Reels mostrando o (produto) numa peça de arte em vidro, no cenário de obra]. Meus Reels já bateram 5,5M a 8M de views.\n"
"\n"
"Anexei o media kit completo. Topa uma call de 15 min essa semana pra eu te mostrar o formato?\n"
"\n"
"Abraço,\n"
"Augusto Felipe · @fidevidraceiro\n"
"augustofelipe@futurah.co · WhatsApp +55 34 9254-8114")

h2("Template D — Follow-up (3–5 dias depois, sem resposta)", "D")
template_box(
"Oi [nome], passando rapidinho pra não deixar a ideia se perder. 🙂\n"
"\n"
"Sei que a rotina é corrida — se fizer sentido, consigo adaptar o formato ao que a [marca] está priorizando agora. Quer que eu mande uma versão mais enxuta da ideia?")

# =====================================================================
# 5. RATE CARD
# =====================================================================
doc.add_page_break()
kicker("Manual de venda")
h1("Da resposta ao fechamento")
body("O que fazer quando a marca responde. Esta é a parte que decide o dinheiro: preço, objeção, call, contrato e retenção.", GRAY)

h2("Rate card 2026 (referência interna)", "5")
body("Base: a proposta Augusto × BYD. Valores de partida — ajuste por marca, esforço e uso. Regra nº 1: NUNCA dê preço no DM. Preço só na call ou no e-mail, depois de entender o objetivo da marca.")

h3("Pacotes base — Ativação Digital")
ref_table("Pacote", "Valor de referência", [
    ("3 Stories", "R$ 2.000"),
    ("1 Reels em collab", "R$ 8.000"),
    ("Combo (Reels + Stories)", "R$ 8.000  (de R$ 10.000)"),
], left_w=2.6)

h3("Ativações premium — Art Experiences")
small("Arte feita sobre o produto/marca — cada nível, uma tela maior. Todas acompanham Reels + Stories. É o diferencial que só o Fi de Vidraceiro entrega.")
ref_table("Nível", "Valor de referência", [
    ("Art Experience — símbolo/logo da marca em vidro (~50cm)", "R$ 20.000"),
    ("Wheel / peça artística personalizada", "R$ 30.000"),
    ("Full Art Experience — produto inteiro virando obra", "R$ 85.000"),
], left_w=3.9)

h3("Regras comerciais")
bullet("Valores incluem produção, gravação, postagem e operação completa.", "O que está no preço: ")
bullet("impulsionamento/mídia paga, whitelisting, direitos de uso estendidos (30/60/90 dias) e exclusividade de categoria. Sempre cobrados à parte.", "Fora do preço (add-ons): ")
bullet("cross-post no TikTok (430k) e Kwai (110k) como upsell — o mesmo Reels alcançando +39M/mês.", "Multiplataforma: ")
bullet("50% antecipado (produção) + 50% na entrega.", "Pagamento: ")
bullet("só se o produto tiver valor alto E a marca for estratégica pra virar case. Senão: permuta + cachê reduzido. Nunca permuta pura por marca fraca.", "Permuta: ")

callout("Esqueleto da proposta (base BYD) — pra reaproveitar em qualquer marca",
        "Capa → Quem é o Augusto (bio) → Mídia Kit (números + público) → Marcas que confiam → Prova (Reels de maior alcance) → "
        "Ativação Digital (pacotes base) → Art Experiences (escada de valor) → Condições comerciais → Fechamento com CTA. "
        "Troca o contexto do produto e os renders; a estrutura não muda.", "F5F5F5", DARK)

# =====================================================================
# 6. OBJEÇÕES
# =====================================================================
h2("Respostas às objeções mais comuns", "6")
body("Improvisar aqui mata deal. Tenha a resposta pronta — e SEMPRE termine amarrando um próximo passo com data.")
ref_table("A marca diz…", "Você responde…", [
    ("\"Está caro\"", "\"O valor reflete produção completa + alcance de 59M/mês. Consigo montar um pacote menor pra começar (ex.: 1 Reels) e a gente escala com o resultado.\""),
    ("\"Não temos budget agora\"", "\"Sem problema. Quando vocês fecham o planejamento do próximo trimestre? Te procuro em [mês].\" → registrar na planilha e cumprir."),
    ("\"Topamos permuta / produto\"", "\"Permuta funciona quando o produto tem valor alto e a marca é estratégica pra virar case. Consigo permuta + um cachê reduzido pra viabilizar a produção.\""),
    ("\"Manda o media kit que a gente avalia\"", "Mandar na hora, MAS amarrar: \"Mando agora — te chamo quinta pra ouvir a impressão de vocês.\" Nunca deixar em aberto."),
    ("\"Já trabalhamos com outro criador\"", "\"Ótimo, quer dizer que vocês já veem valor em criador. Meu diferencial é transformar o produto de vocês em arte de vidro — um formato que nenhum outro criador entrega. Complementa, não compete.\""),
], left_w=2.4)

# =====================================================================
# 7. CALL
# =====================================================================
h2("Script da call de 15 minutos", "7")
body("O momento de maior alavancagem do funil. Regra de ouro: sair da call com um prazo combinado — nunca com \"a gente se fala\".")
ref_table("Bloco", "O que fazer", [
    ("0–3 min · Ouvir", "Perguntar o objetivo da marca no trimestre: awareness? lançamento? empurrar varejo? Deixar a marca falar."),
    ("3–8 min · A ideia", "Apresentar UMA ideia adaptada ao que você acabou de ouvir + mostrar o case do Reel de 8M como prova."),
    ("8–12 min · Preço", "Apresentar os pacotes e a faixa de valor. Ancorar no premium (Art Experience) e oferecer o pacote base como entrada."),
    ("12–15 min · Fechar", "Combinar o próximo passo COM DATA: \"Mando a proposta fechada até sexta e a gente conversa segunda.\""),
], left_w=2.2)

# =====================================================================
# 8. DEPOIS DO SIM
# =====================================================================
h2("Depois do \"sim\": o operacional que protege o deal", "8")
body("Um \"sim\" mal amarrado vira calote, retrabalho infinito ou dor de cabeça jurídica. Checklist mínimo:")
bullet("escopo, nº de entregas, prazo de veiculação, forma de pagamento (50/50) e prazo de emissão da NF.", "Contrato simples: ")
bullet("no máximo 1–2 rodadas de alteração de roteiro. Deixar isso escrito evita revisão sem fim.", "Alterações: ")
bullet("por quanto tempo a marca pode usar o conteúdo (30/60/90 dias). Uso em mídia paga é cobrado à parte.", "Direitos de uso: ")
bullet("de categoria, com prazo e preço definidos. Não impedir concorrentes de graça.", "Exclusividade: ")
bullet("toda publi identificada — #publi + a ferramenta de parceria paga do Instagram. É regra do CONAR e marca séria vai exigir.", "Transparência: ")
bullet("briefing alinhado e roteiro aprovado ANTES de gravar. Combinar link/cupom rastreável para medir resultado.", "Antes de gravar: ")

# =====================================================================
# 9. RELATÓRIO + RETENÇÃO
# =====================================================================
h2("Relatório pós-campanha & retenção", "9")
body("Caçar marca nova é caro; renovar é onde está o dinheiro (e o inbound). Prova de ROI é o que transforma um deal em recorrência. Envie um mini-relatório ~7 dias após a publi:")
ref_table("Métrica", "O que mostrar", [
    ("Alcance & Views", "Total de contas alcançadas e visualizações, por plataforma."),
    ("Engajamento", "Curtidas, comentários, salvamentos e compartilhamentos + retenção do vídeo."),
    ("Comentários qualitativos", "Prints dos melhores comentários do público sobre a marca/produto."),
    ("Conversão", "Cliques no link ou usos do cupom rastreável (combinado ANTES da campanha)."),
    ("Comparativo", "Como esses números se saem vs. a média do perfil — mostrar que performou."),
], left_w=2.4)
callout("O gancho de renovação",
        "Feche o relatório com a proposta de continuidade: \"Com esse resultado, faz sentido a gente montar um pacote trimestral?\" "
        "É o caminho mais barato pro próximo deal.", "EFFaF0", GREEN)

# =====================================================================
# 10. AGÊNCIA VS MARCA
# =====================================================================
h2("Agência vs. marca direta (+ inbound passivo)", "10")
body("Nas marcas grandes da planilha (Coca-Cola, McDonald's, Nubank), a porta real costuma ser uma agência de influência ou plataforma, não o marketing interno. Muda a abordagem:")
ref_table("Canal", "Como se comporta", [
    ("Agência de influência", "Decide por CPM/entregável, responde mais rápido — mas espreme preço. Seja objetivo com números e pacotes."),
    ("Marca direta", "Valoriza a ideia criativa e a conexão com o público. Ciclo mais longo, mas ticket e relação melhores."),
], left_w=2.4)
body("Para gerar inbound passivo (a marca vem até você), cadastre o perfil do Augusto nas plataformas de creator do Brasil — ex.: Squid, Influency.me e BrandLovers.", GRAY)

# =====================================================================
# 11. PLANILHA DE ALVOS
# =====================================================================
from docx.enum.section import WD_ORIENT, WD_SECTION
_sec = doc.add_section(WD_SECTION.NEW_PAGE)
_sec.orientation = WD_ORIENT.LANDSCAPE
_sec.page_width, _sec.page_height = _sec.page_height, _sec.page_width
_sec.left_margin = Inches(0.6); _sec.right_margin = Inches(0.6)
_sec.top_margin = Inches(0.6); _sec.bottom_margin = Inches(0.6)

# ---- Diretório de contatos (tabela de 4 colunas em paisagem) ----
def contacts_table(rows):
    cols = ["Marca", "Contato", "Cargo / papel", "Melhor canal (confirme antes de enviar)"]
    t = doc.add_table(rows=1, cols=len(cols))
    t.style = "Table Grid"; t.autofit = False
    hdr = t.rows[0].cells
    for j,c in enumerate(cols):
        hdr[j].text = c; shade(hdr[j], "1A1A1A"); set_cell_font(hdr[j], 8.5, True, white=True)
    widths = [1.3, 2.1, 2.5, 3.6]
    last = None
    for (marca, contato, cargo, canal, star) in rows:
        cells = t.add_row().cells
        show = marca if marca != last else ""
        last = marca
        cells[0].text = show
        cells[1].text = ("★ " if star else "") + contato
        cells[2].text = cargo
        cells[3].text = canal
        for j in range(4):
            set_cell_font(cells[j], 8, bold=(j==0 or (j==1 and star)))
        if show:
            shade(cells[0], "EAF4FF")
        if star:
            shade(cells[1], "EFFaF0")
    for row in t.rows:
        for j,w in enumerate(widths):
            row.cells[j].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t

h2("Diretório de contatos por marca", "11")
body("Várias opções por marca — se um não responde, você tem alternativa. ★ marca a melhor porta de entrada. Em bebida, Coca-Cola, fast food e Havaianas, a AGÊNCIA de influência costuma responder mais rápido que o time interno. Cargos giram: confirme no LinkedIn antes de abordar.")
contacts_table([
    # CIMED
    ("Cimed", "Toguro", "Head de Comunicação — é criador (11M)", "IG @toguro · toguro@mansaomaromba.com", True),
    ("Cimed", "Camille Lau", "CMO (desde out/25)", "LinkedIn: buscar \"Camille Lau Cimed\"", False),
    ("Cimed", "Gabriela Masson", "Gerente de Marketing Sênior", "LinkedIn /gabrielamasson", False),
    ("Cimed", "Karla Felmanas", "Sócia (família controladora)", "IG @karlacimed", False),
    ("Cimed", "Agência (criação)", "\"Cimed Squad\" — BFerraz lidera", "Buscar atendimento Cimed na BFerraz", False),
    # STANLEY
    ("Stanley", "Patrícia Esteves", "Dir. Executiva de Mkt LatAm (CMO)", "LinkedIn /paesteves", True),
    ("Stanley", "Babi (Bárbara) Coelho", "Ger. de Mkt e Comunicação BR", "Buscar \"Babi Coelho Stanley 1913\"", False),
    ("Stanley", "Chloe Sestero", "Coordenadora de Marketing", "Buscar \"Chloe Sestero Stanley\"", False),
    ("Stanley", "Programa próprio", "BearForce (embaixadores) + Awin", "stanley1913.com.br/BearForce", False),
    # NATURA
    ("Natura", "Tatiana Ponce", "CMO Natura&Co", "Buscar \"Tatiana Ponce Natura\"", True),
    ("Natura", "Nathalia Beltrame", "Gerente de Marketing", "LinkedIn: buscar \"Nathalia Beltrame Natura\"", False),
    ("Natura", "Prog. Criadores da Beleza", "operado por agências", "Contato: YouPix + Flint", False),
    ("Natura", "Gabriella Cartaxo", "Creator insider (indica quem decide)", "LinkedIn /gabriellacartaxo", False),
    # BOTICÁRIO
    ("O Boticário", "Carolina Carrasco", "Dir. Branding/Comunicação Boticário", "Buscar \"Carolina Carrasco Boticário\"", True),
    ("O Boticário", "Renata Gomide", "VP de Marketing (Grupo)", "IG @reagomide", False),
    ("O Boticário", "Marcela De Masi", "Dir. Exec. Branding/Comunicação", "LinkedIn /marcelademasi", False),
    ("O Boticário", "Agência (influência)", "CoCreators atende a marca", "Andrea Iorio IG @aiorio_br", False),
    # COCA-COLA
    ("Coca-Cola", "Raquel Ribeiro", "Head de Marketing Brasil (out/25)", "Buscar \"Raquel Ribeiro Coca-Cola\"", True),
    ("Coca-Cola", "Mariana David", "Senior Brand Director Brasil", "LinkedIn /marianadavid", False),
    ("Coca-Cola", "Raphael Taborda", "Social Media / CRM (influência operacional)", "LinkedIn /raphael-taborda0001", False),
    ("Coca-Cola", "Agência (influência)", "BR Media Group (hub) · Spark", "Buscar conta Coca na BR Media Group", False),
    # RED BULL
    ("Red Bull", "Culture Marketing Manager", "quem trabalha com artistas/criadores", "Buscar \"Culture Marketing Red Bull Brasil\"", True),
    ("Red Bull", "Diego Colicchio", "MD & Presidente Brasil", "LinkedIn /diegocolicchio", False),
    ("Red Bull", "Agência", "Publicis (criação/mídia) · Ágora (PR)", "Via Publicis", False),
    # HEINEKEN
    ("Heineken", "Fernanda Saboya", "Dir. Consumer Connections (concentra influência)", "LinkedIn /fsaboya", True),
    ("Heineken", "Igor de Castro", "Dir. Comunicação e Branding (Heineken)", "LinkedIn /igordecastrooliveira", False),
    ("Heineken", "Carolina Monteiro", "Brand Coordinator/Manager", "Buscar \"Carolina Monteiro Heineken\"", False),
    ("Heineken", "Agência (influência)", "BR Media/StarLab (hub) · Live (Amstel)", "Via agência, por marca", False),
    # AMBEV
    ("Ambev", "Mariana Dedivitis", "Marketing Director — Corona", "LinkedIn /mariana-dedivitis", True),
    ("Ambev", "Guilherme Almeida", "Execution Director — Brahma", "LinkedIn /guilhermealmeidaaguiarferreira", False),
    ("Ambev", "Leandro Mendonça", "Dir. Patrocínios e Experiências", "Buscar no LinkedIn (bom p/ ativação)", False),
    ("Ambev", "Agência (influência)", "Draftline/Soko (in-house) · InPress (PR)", "Marcelo Cia / Thais Szpigel (InPress)", False),
    # NUBANK
    ("Nubank", "Raquel Forastieri", "Brand Mkt Manager (faz influencer mgmt)", "LinkedIn /rquelforastieri", True),
    ("Nubank", "Juliana Roschel", "CMO Brasil", "Buscar \"Juliana Roschel Nubank\"", False),
    ("Nubank", "Diego Nunes", "Senior Marketing Manager", "LinkedIn /diegoanunes", False),
    ("Nubank", "(sem agência)", "influência é 100% in-house", "Ir pelo time de Brand Marketing", False),
    # IFOOD
    ("iFood", "Marina Moreira", "Head de Social & Influência — melhor alvo da lista", "LinkedIn /marina-moreira-97520696", True),
    ("iFood", "Renata Lamarco", "Diretora de Marketing (escopo inclui influência)", "LinkedIn /renatalamarco", False),
    ("iFood", "Vivian Jung", "Head de Branding / Brand Experience", "Buscar \"Vivian Jung iFood\"", False),
    ("iFood", "Ana Gabriela Lopes", "CMO (acima de todas)", "LinkedIn /anagabrielaoliveiralopes", False),
    # HAVAIANAS
    ("Havaianas", "Karen Bartels", "Sócia/MD da a.gente (agência de influência)", "Buscar \"Karen Bartels a.gente\"", True),
    ("Havaianas", "Juliana Soncini", "Diretora de Marketing LATAM", "LinkedIn: buscar \"Juliana Soncini Havaianas\"", False),
    ("Havaianas", "Fernanda Botteghin", "Coord. de Comunicação Havaianas", "LinkedIn /fernanda-botteghin", False),
    ("Havaianas", "Agência (digital+influência)", "GALERIA.ag assumiu a conta", "galeria.ag", False),
    # SAMSUNG (celular)
    ("Samsung (celular)", "gen_c — Cristiane Vieira", "Dir. Atendimento da agência de influência Mobile", "Buscar \"Cristiane Vieira gen_c\"", True),
    ("Samsung (celular)", "Mario Sousa", "Dir. Sênior de Mkt Mobile", "Buscar \"Mario Sousa Samsung\"", False),
    ("Samsung (celular)", "Rodrigo S. Menezes", "Mkt Digital Mobile (#TeamGalaxy)", "Buscar no LinkedIn", False),
    ("Samsung (celular)", "Hugo Harten", "Ger. Mkt Brand Experience Mobile", "LinkedIn /hugo-harten", False),
    # MOTOROLA
    ("Motorola", "Andrea Brandi", "Sr. Brand Mkt Manager Brasil", "LinkedIn /andrea-brandi", True),
    ("Motorola", "Stella Colucci", "Diretora de Marketing Brasil", "Buscar \"Stella Colucci Motorola\"", False),
    ("Motorola", "Alessandra Barcala", "CMO América Latina", "LinkedIn /alessandrabarcala", False),
    # MCDONALD'S
    ("McDonald's", "Guilherme Coe", "Gerente de Marketing (confirmado 2025)", "Buscar \"Guilherme Coe McDonald's\"", True),
    ("McDonald's", "Ilca Sierra", "CMO Brasil (mar/25)", "Buscar \"Ilca Sierra McDonald's\"", False),
    ("McDonald's", "(cargo a buscar)", "Branded Content / Social", "Buscar \"Social Arcos Dorados\" no LinkedIn", False),
    # BURGER KING
    ("Burger King", "Aymê Alves", "Coord. Comunicação e Redes Sociais (Zamp)", "Buscar \"Aymê Alves Zamp\"", True),
    ("Burger King", "Pedro Barbosa", "CMO BK (jan/26)", "Buscar \"Pedro Barbosa Zamp\"", False),
    ("Burger King", "Igor Puga", "VP de Marketing Zamp", "Buscar \"Igor Puga Zamp\"", False),
    # GWM
    ("GWM", "Lucas Coelho", "CMO / Diretor de Marketing", "LinkedIn /lucas-a-coelho", True),
    ("GWM", "André Leite", "Dir. de Marketing e Produto", "Buscar \"André Leite GWM\"", False),
    ("GWM", "Magda Lima", "Gerente de Marketing", "LinkedIn /magdarlima", False),
    ("GWM", "Ricardo Bastos", "Dir. Relações Institucionais (muito público)", "Buscar \"Ricardo Bastos GWM\"", False),
    # CAOA CHERY
    ("Caoa Chery", "Jan Telecki", "Diretor de Marketing (grupo CAOA)", "Buscar \"Jan Telecki CAOA\"", True),
    ("Caoa Chery", "Tai Kawasaki", "Ger. Executivo Comercial (mkt + vendas)", "Buscar no LinkedIn", False),
    ("Caoa Chery", "Agência", "Baila (full-service) · Z515", "Buscar atendimento Caoa na Baila", False),
])
small("⚠ Já saíram do cargo (NÃO usar): Ted Ketterer (Coca-Cola) · Sérgio Eleutério (McDonald's) · Mariana Rhormens e Nathalia Pires (Havaianas) · provavelmente Fernando Sodré (Cimed). @ de executivo raramente é público — onde não há IG, o caminho é LinkedIn. Fontes: Meio & Mensagem, propmark, Exame, Forbes, LinkedIn (2024-2026).")

doc.add_page_break()
h2("Planilha de alvos & follow-up", "12")
body("Copie esta tabela pro Excel/Sheets e trabalhe uma linha por marca. Status sugerido: Aquecendo → DM enviado → Respondeu → Media kit enviado → Call → Fechado / Descartado.")

cols = ["Marca", "Decisor / cargo", "Contato decisor", "Canal", "Data DM", "Aqueci?", "Status", "Próx. follow-up", "Valor R$", "Notas"]
tblp = doc.add_table(rows=1, cols=len(cols))
tblp.style = "Table Grid"
for j,c in enumerate(cols):
    cell = tblp.rows[0].cells[j]
    cell.text = c
    shade(cell, "1A1A1A"); set_cell_font(cell, 8.5, True, white=True)
# 12 marcas-alvo sugeridas — qualquer nicho, filtro = hype + verba de influenciador
examples = [
    ["Cimed", "Toguro (Head Com.) / Camille Lau (CMO)", "IG + LinkedIn", "", "Não", "A aquecer", "", "Farma viral; Toguro é criador, melhor porta. Ideia: Carmed/produto gigante em vidro."],
    ["Stanley", "Babi Coelho (Ger. Mkt) — confirmar", "IG + LinkedIn", "", "Não", "A aquecer", "", "Hype absurdo, público jovem. Ideia: o copo virando obra de vidro."],
    ["Natura / Boticário", "C. Carrasco (Botic.) / prog. Criadores da Beleza", "IG + LinkedIn", "", "Não", "A aquecer", "", "Perfume É vidro. Natura via agência YouPix. Frasco gigante esculpido."],
    ["Coca-Cola", "Raquel Ribeiro (Head Mkt BR)", "LinkedIn / Agência", "", "Não", "A aquecer", "", "Verba enorme; publi via agência. A garrafa contour em vidro. Marca-sonho."],
    ["Red Bull", "via agência Publicis / Culture Mkt", "LinkedIn / Agência", "", "Não", "A aquecer", "", "Sem nome público; conta na Publicis. Asas/lata em vidro."],
    ["Heineken / Ambev", "F. Saboya (Heineken) / ag. InPress (Ambev)", "LinkedIn / Agência", "", "Não", "A aquecer", "", "Rasga verba. Heineken tem dona de influência nominal. Garrafa/estrela em vidro."],
    ["Nubank", "Raquel Forastieri (Brand Mkt, faz influência)", "IG + LinkedIn", "", "Não", "A aquecer", "", "Influência in-house. CMO: Juliana Roschel. Peça roxa / cartão gigante."],
    ["iFood", "Marina Moreira (Head Social & Influência)", "LinkedIn", "", "Não", "A aquecer", "", "Melhor porta da lista. /marina-moreira-97520696. Ativação vermelha."],
    ["Havaianas", "Karen Bartels (ag. a.gente) / J. Soncini", "LinkedIn / Agência", "", "Não", "A aquecer", "", "Influência via GALERIA.ag. Chinelo/verão em arte de vidro."],
    ["Samsung / Motorola", "ag. gen_c (Samsung mobile) / Andrea Brandi (Moto)", "LinkedIn / Agência", "", "Não", "A aquecer", "", "gen_c cuida da influência Mobile. Mirar linha CELULAR."],
    ["McDonald's / BK", "Guilherme Coe (McD) / Aymê Alves (BK redes)", "LinkedIn / Agência", "", "Não", "A aquecer", "", "QSR investe muito em criador. Batata/casquinha em vidro."],
    ["GWM / Caoa Chery", "Lucas Coelho (GWM) / Jan Telecki (Caoa)", "LinkedIn", "", "Não", "A aquecer", "", "Reaproveita o case BYD: carro virando obra. Lucas tem LinkedIn público."],
]
# realinhar pras 10 colunas (insere "Contato decisor" no índice 2 e "Valor R$" antes de Notas)
examples = [[r[0], r[1], "", r[2], r[3], r[4], r[5], r[6], "", r[7]] for r in examples]
for _ in range(3):
    examples.append([""]*len(cols))
for row in examples:
    cells = tblp.add_row().cells
    for j,val in enumerate(row):
        cells[j].text = val
        set_cell_font(cells[j], 8.5, color=DARK)

small("Meta prática: 10–15 marcas aquecendo por vez · 5–10 DMs novos/dia, só pra marcas já aquecidas · espaçar envios (2–5 min).")

# =====================================================================
# 13. DMs PERSONALIZADAS (volta pra retrato)
# =====================================================================
_sec3 = doc.add_section(WD_SECTION.NEW_PAGE)
_sec3.orientation = WD_ORIENT.PORTRAIT
_sec3.page_width = Inches(8.5); _sec3.page_height = Inches(11)
_sec3.left_margin = Inches(1); _sec3.right_margin = Inches(1)
_sec3.top_margin = Inches(1); _sec3.bottom_margin = Inches(1)

def meta_box(contato, cargo, canal, ideia):
    tbl = doc.add_table(rows=1, cols=1); cell = tbl.rows[0].cells[0]
    shade(cell, "EAF4FF")
    tcPr = cell._tc.get_or_add_tcPr(); tcMar = OxmlElement("w:tcMar")
    for tag in ("top","start","bottom","end"):
        m = OxmlElement(f"w:{tag}"); m.set(qn("w:w"), "120"); m.set(qn("w:type"), "dxa"); tcMar.append(m)
    tcPr.append(tcMar)
    for i,(lab,val) in enumerate([("Contato:", f"{contato} — {cargo}"), ("Canal:", canal), ("Ideia de arte:", ideia)]):
        p = cell.paragraphs[0] if i==0 else cell.add_paragraph()
        rl = p.add_run(lab + " "); rl.font.bold = True; rl.font.size = Pt(9.5); rl.font.color.rgb = ACCENT
        rv = p.add_run(val); rv.font.size = Pt(9.5)
        p.paragraph_format.space_after = Pt(1)
    return tbl

def dm_box(text):
    tbl = doc.add_table(rows=1, cols=1); cell = tbl.rows[0].cells[0]
    shade(cell, "F5F5F5")
    tcPr = cell._tc.get_or_add_tcPr(); borders = OxmlElement("w:tcBorders")
    left = OxmlElement("w:left"); left.set(qn("w:val"),"single"); left.set(qn("w:sz"),"18")
    left.set(qn("w:space"),"0"); left.set(qn("w:color"),"1EA05A"); borders.append(left); tcPr.append(borders)
    tcMar = OxmlElement("w:tcMar")
    for tag in ("top","start","bottom","end"):
        m = OxmlElement(f"w:{tag}")
        m.set(qn("w:w"), "160" if tag in ("start","end") else "120"); m.set(qn("w:type"), "dxa"); tcMar.append(m)
    tcPr.append(tcMar)
    for i,line in enumerate(text.split("\n")):
        p = cell.paragraphs[0] if i==0 else cell.add_paragraph()
        r = p.add_run(line); r.font.size = Pt(10.5); r.font.name = "Calibri"
        p.paragraph_format.space_after = Pt(2)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return tbl

def entry(num, marca, contato, cargo, canal, ideia, dm):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(4)
    r0 = p.add_run(f"{num:02d}  "); r0.font.size = Pt(14); r0.font.bold = True; r0.font.color.rgb = ACCENT
    r = p.add_run(marca); r.font.size = Pt(14); r.font.bold = True; r.font.color.rgb = DARK
    meta_box(contato, cargo, canal, ideia)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    dm_box(dm)

kicker("Parte prática · DMs prontas")
h1("DMs personalizadas por marca")
body("Uma mensagem por marca (a melhor porta de entrada). Modelo INTEREST, < 150 palavras. Personalize o [colchete] em cada envio.", GRAY)
h3("Antes de enviar — 5 regras")
for t in [
    "Personalize o [colchete]: cite algo que a pessoa/marca postou nas ÚLTIMAS 2 SEMANAS.",
    "NÃO mande proposta nem preço no 1º contato — o objetivo é abrir conversa e achar o decisor.",
    "Melhor janela: terça a quinta, 10h–14h. Um CTA só.",
    "Onde há @ do Instagram, mande DM pelo @fidevidraceiro. Onde só há LinkedIn, use a nota do convite.",
    "Confirme o cargo atual no LinkedIn antes de enviar — cargos de marketing giram rápido.",
]:
    pp = doc.add_paragraph(style="List Bullet"); pp.add_run(t).font.size = Pt(10)

entry(1, "Cimed", "Toguro", "Head de Comunicação (é criador, 11M)",
      "Instagram: @toguro  ·  e-mail na bio: toguro@mansaomaromba.com  ·  alt: @karlacimed (sócia)",
      "Carmed / produto do momento recriado gigante, esculpido em vidro.",
"""Fala Toguro! Aqui é o Augusto, o Fi de Vidraceiro (@fidevidraceiro).

Acompanho a treta boa que a Cimed faz no marketing — [cite a última ação/vídeo de vocês]. Eu transformo vidro em arte pra uma comunidade de 1.2M (59M de alcance/mês) e tô com uma ideia que tem a sua cara: recriar o [Carmed / produto do momento] gigante, esculpido em vidro, virando um Reels de bater milhão. Meus Reels já passaram de 5,5M a 8M de views.

Você toca a comunicação aí — topa eu te mandar o media kit e o conceito? 🙏""")

entry(2, "Stanley", "Patrícia Esteves", "Diretora Executiva de Marketing LatAm",
      "LinkedIn: /paesteves  ·  (IG não público — buscar \"Patrícia Esteves Stanley\")",
      "O copo Stanley recriado como obra de vidro feita à mão.",
"""Oi Patrícia, tudo bem? Sou o Augusto Felipe, criador do Fi de Vidraceiro (@fidevidraceiro — 1.2M seguidores, 59M de alcance/mês). Transformo produtos em arte de vidro.

Vi o quanto o copo Stanley virou objeto de desejo no Brasil [cite ação/lançamento recente] e tive uma ideia: recriar o copo como uma obra de vidro feita à mão, num Reels de bastidor. Marcas como Spotify, Bombril e TekBond já toparam meu trabalho.

Quem cuida das collabs com criador aí? Posso mandar meu media kit. Abraço!""")

entry(3, "Natura", "Tatiana Ponce", "CMO Natura&Co",
      "LinkedIn: buscar \"Tatiana Ponce Natura\"  ·  alt: programa Criadores da Beleza (via YouPix)",
      "Frasco de perfume gigante esculpido em vidro (perfume É vidro).",
"""Oi Tatiana! Sou o Augusto, o Fi de Vidraceiro (@fidevidraceiro — 1.2M seguidores, 59M/mês). Meu diferencial é transformar produto em arte de vidro — e perfume é vidro, então o encaixe com a Natura é natural.

Ideia: esculpir um frasco gigante de [produto/linha] em vidro, virando um Reels. Meus Reels já bateram 5,5M–8M de views.

Quem cuida das parcerias com criador / do programa Criadores da Beleza? Adoraria mandar o media kit.""")

entry(4, "O Boticário", "Renata Gomide", "VP de Marketing (Grupo Boticário)",
      "Instagram: @reagomide  ·  operacional: Carolina Carrasco (Dir. Branding) / ag. CoCreators (@aiorio_br)",
      "Frasco de perfume do Boticário gigante, esculpido em vidro.",
"""Oi Renata! Sou o Augusto, o Fi de Vidraceiro (@fidevidraceiro — 1.2M seguidores, 59M/mês). Transformo produto em arte de vidro.

Perfume é vidro — dá pra fazer um frasco de [linha/lançamento] gigante esculpido, virando um Reels lindo. [cite o lançamento recente de vocês]. Meus Reels já passaram de 8M de views.

Quem toca as collabs com criador no Boticário? Posso te mandar o media kit.""")

entry(5, "Coca-Cola", "Raquel Ribeiro", "Head de Marketing Brasil",
      "LinkedIn: buscar \"Raquel Ribeiro Coca-Cola\"  ·  influência via agência BR Media Group",
      "A garrafa contour recriada à mão em vidro.",
"""Oi Raquel, tudo bem? Sou o Augusto Felipe, o Fi de Vidraceiro (@fidevidraceiro — 1.2M seguidores, 59M de alcance/mês). Transformo produtos em obra de arte de vidro.

A garrafa contour é o objeto mais icônico do mundo — imagina ela recriada à mão em vidro, num Reels. [cite campanha recente da Coca]. Marcas como Spotify e Bombril já confiaram no trabalho.

Quem cuida das collabs com criador aí (ou qual agência)? Mando o media kit.""")

entry(6, "Red Bull", "Culture Marketing Manager", "quem lida com artistas/criadores",
      "LinkedIn: buscar \"Culture Marketing Red Bull Brasil\"  ·  ou via agência Publicis",
      "Obra de vidro com as asas / a lata da Red Bull.",
"""Oi [nome], sou o Augusto Felipe, o Fi de Vidraceiro (@fidevidraceiro — 1.2M seguidores). A Red Bull patrocina o inusitado — e eu faço arte de vidro que vira Reels de milhões (já bati 8M).

Ideia: uma obra de vidro com as asas / a lata da Red Bull, no maior estilo "te dá asas". [cite um evento/ativação recente de vocês].

Você cuida de Culture / creators aí? Bora trocar ideia — mando o media kit.""")

entry(7, "Heineken", "Fernanda Saboya", "Dir. Consumer Connections (concentra influência)",
      "LinkedIn: /fsaboya  ·  (é board member da BCMA — branded content)",
      "Garrafa verde / estrela Heineken esculpida em vidro.",
"""Oi Fernanda! Sou o Augusto, o Fi de Vidraceiro (@fidevidraceiro — 1.2M seguidores, 59M/mês). Vi que você toca a frente de influência e branded content da Heineken.

Transformo produto em arte de vidro — dá pra esculpir a garrafa verde / a estrela num Reels lindo pra [ocasião / campanha]. Meus Reels já bateram 5,5M–8M de views.

Topa eu te mandar o media kit e o conceito?""")

entry(8, "Ambev (Corona)", "Mariana Dedivitis", "Marketing Director — Corona",
      "LinkedIn: /mariana-dedivitis  ·  influência via Draftline/Soko · InPress (PR)",
      "A garrafa da Corona recriada em vidro num pôr do sol.",
"""Oi Mariana! Sou o Augusto, o Fi de Vidraceiro (@fidevidraceiro — 1.2M seguidores). Faço arte de vidro que vira Reels de milhões (já bati 8M).

A Corona tem a garrafa mais fotogênica do mundo — imagina ela recriada em vidro num pôr do sol, virando conteúdo. [cite verão / campanha recente da Corona].

Quem cuida das collabs com criador da Corona? Mando o media kit. Saúde!""")

entry(9, "Nubank", "Raquel Forastieri", "Brand Marketing Manager (cuida de influência)",
      "LinkedIn: /rquelforastieri  ·  influência é 100% in-house (sem agência)",
      "Peça roxa esculpida — o cartão / o 'Nu' virando obra de vidro.",
"""Oi Raquel! Sou o Augusto, o Fi de Vidraceiro (@fidevidraceiro — 1.2M seguidores, 59M/mês). Transformo produto em arte de vidro.

Ideia pro Nubank: uma peça roxa esculpida — o cartão / o "Nu" virando obra num Reels. Vi que você cuida da frente de influência aí. Meus Reels já passaram de 8M de views.

Posso te mandar o media kit?""")

entry(10, "iFood", "Marina Moreira", "Head de Social & Influência (melhor porta da lista)",
      "LinkedIn: /marina-moreira-97520696  ·  acima: Renata Lamarco (Dir. Mkt)",
      "Ativação divertida em vidro com a identidade vermelha do iFood.",
"""Oi Marina! Sou o Augusto Felipe, o Fi de Vidraceiro (@fidevidraceiro — 1.2M seguidores, 59M de alcance/mês). Vi que você lidera social & influência no iFood.

Transformo produto em arte de vidro — dá pra criar uma ativação divertida em vidro com a identidade do iFood, virando um Reels de milhões (já bati 8M). [cite uma ação recente de vocês].

Topa eu te mandar o media kit com o conceito?""")

entry(11, "Havaianas", "Karen Bartels", "Sócia/MD da a.gente (agência de influência da marca)",
      "LinkedIn: buscar \"Karen Bartels a.gente\"  ·  a.gente é o braço de influência da GALERIA.ag",
      "Chinelo Havaianas traduzido numa obra de vidro de verão.",
"""Oi Karen! Sou o Augusto, o Fi de Vidraceiro (@fidevidraceiro — 1.2M seguidores). Sei que a a.gente toca a influência da Havaianas.

Meu diferencial é transformar produto em arte de vidro — imagina um chinelo Havaianas traduzido numa obra de vidro de verão, virando um Reels. Já passei de 8M de views.

Vocês estão montando ações de criador pra Havaianas? Mando meu media kit.""")

entry(12, "Samsung (celular)", "Cristiane Vieira", "Dir. Atendimento da gen_c (agência de influência do Galaxy)",
      "LinkedIn: buscar \"Cristiane Vieira gen_c\"  ·  gen_c cuida do #TeamGalaxy / Mobile",
      "A tela / o aparelho esculpido em vidro (gancho literal: Gorilla Glass).",
"""Oi Cristiane! Sou o Augusto Felipe, o Fi de Vidraceiro (@fidevidraceiro — 1.2M seguidores). Sei que a gen_c cuida da influência do Galaxy.

Faço arte de vidro — e o gancho com celular é literal (Gorilla Glass): dá pra esculpir a tela / o aparelho em vidro, num Reels tech de milhões (já bati 8M).

Vocês estão abrindo pauta de creators pro próximo lançamento? Mando o media kit.""")

entry(13, "Motorola", "Andrea Brandi", "Sr. Brand Marketing Manager Brasil",
      "LinkedIn: /andrea-brandi  ·  acima: Stella Colucci (Dir. Mkt) / Alessandra Barcala (CMO LATAM)",
      "O aparelho Motorola retratado/esculpido em vidro.",
"""Oi Andrea! Sou o Augusto, o Fi de Vidraceiro (@fidevidraceiro — 1.2M seguidores, 59M/mês). Transformo produto em arte de vidro — com celular o gancho é literal (a tela é vidro).

Ideia: o Motorola [modelo] esculpido / retratado em vidro num Reels. Já bati 8M de views.

Quem cuida das collabs com criador aí? Mando o media kit.""")

entry(14, "McDonald's", "Guilherme Coe", "Gerente de Marketing",
      "LinkedIn: buscar \"Guilherme Coe McDonald's\"  ·  CMO: Ilca Sierra",
      "A batata / casquinha recriada como obra de vidro.",
"""Oi Guilherme! Sou o Augusto, o Fi de Vidraceiro (@fidevidraceiro — 1.2M seguidores). Faço arte de vidro que vira Reels de milhões (já bati 8M).

Ideia pro Méqui: a batata / casquinha recriada como obra de vidro, num conteúdo divertido. [cite campanha recente do McDonald's].

Quem cuida das collabs com criador aí? Mando o media kit.""")

entry(15, "Burger King", "Aymê Alves", "Coord. Comunicação e Redes Sociais (Zamp)",
      "LinkedIn: buscar \"Aymê Alves Zamp\"  ·  CMO: Pedro Barbosa · VP: Igor Puga",
      "O Whopper / a coroa recriado em arte de vidro.",
"""Oi Aymê! Sou o Augusto, o Fi de Vidraceiro (@fidevidraceiro — 1.2M seguidores, 59M/mês). Vi que você toca as redes do BK.

Transformo produto em arte de vidro — dá pra fazer o Whopper / a coroa em vidro num Reels de bater milhão (já passei de 8M). [cite uma sacada recente de vocês].

Topa eu mandar o media kit com a ideia?""")

entry(16, "GWM", "Lucas Coelho", "CMO / Diretor de Marketing",
      "LinkedIn: /lucas-a-coelho  ·  alt: André Leite (Dir. Mkt e Produto) / Magda Lima (Ger. Mkt)",
      "O carro inteiro virando obra de arte (case pronto: BYD).",
"""Oi Lucas! Sou o Augusto Felipe, o Fi de Vidraceiro (@fidevidraceiro — 1.2M seguidores, 59M de alcance/mês).

Já fiz uma ativação transformando um carro inteiro em obra de arte (case BYD) — Reels de milhões. Adoraria trazer isso pra GWM [cite modelo / lançamento recente de vocês].

Quem cuida do marketing de influência aí? Mando o media kit e o case.""")

entry(17, "Caoa Chery", "Jan Telecki", "Diretor de Marketing (grupo CAOA)",
      "LinkedIn: buscar \"Jan Telecki CAOA\"  ·  influência/criação via agência Baila",
      "O modelo da Caoa Chery virando obra de arte (mesmo conceito do case BYD).",
"""Oi Jan! Sou o Augusto, o Fi de Vidraceiro (@fidevidraceiro — 1.2M seguidores). Faço arte de vidro sobre produtos — inclusive transformei um carro inteiro em obra (case BYD, Reels de milhões).

Dá pra fazer o mesmo com o [Tiggo / modelo] da Caoa Chery. [cite lançamento recente].

Quem cuida das collabs com criador aí? Mando o media kit e o case.""")

# rodapé
hr()
p = doc.add_paragraph()
r = p.add_run("Fi de Vidraceiro · @fidevidraceiro · augustofelipe@futurah.co · +55 34 9254-8114")
r.font.size = Pt(8.5); r.font.color.rgb = LIGHT

out = r"E:\futurah\apps\augustofelipe\Kit-Prospeccao-Marcas-Augusto-Felipe.docx"
doc.save(out)
print("SALVO:", out)
